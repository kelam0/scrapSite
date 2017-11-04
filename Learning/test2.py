# -*- coding: utf-8 -*-
"""
Created on Mon Aug 22 23:03:57 2016

@author: Malek
"""
class record:
   empCount = 0
   def __init__(self, a, b, c):
      self.qty = a
      self.id = b
      self.desc = c
      record.empCount += 1

import docx
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR
from docx.enum.text import WD_BREAK

#create docx document
doc = docx.Document()

#create default font
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = docx.shared.Pt(9)

#add a title saying "Hello World!"
run0 = doc.add_paragraph().add_run()
doc.add_paragraph('Hello world!', 'Title')

#add a header saying "Header" of size between 0 and 4 (here size = 2)
doc.add_heading('Header', 2)

#add a paragraph saying "Hello World!"
doc.add_paragraph('Hello world!')

#new page before
doc.paragraphs[0].runs[0].add_break(WD_BREAK.PAGE)

run1 = doc.add_paragraph().add_run()
paraObj1 = doc.add_paragraph('This is a second paragraph on a second page.')

#apply style on paragraph within no extra line between paragraphs
paraObj2 = doc.add_paragraph('first item in unordered list', style = 'List Bullet')


run2 = doc.add_paragraph().add_run()
#special font
style2 = doc.styles['Body Text']
style2.font.name = 'Calibri'
style2.font.size = docx.shared.Pt(11)
style2.font.color.theme_color = MSO_THEME_COLOR.TEXT_1
style2.font.color.theme_color.brightness = 0.4
style2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

run3 = doc.add_paragraph().add_run()
#add a paragraph with special font, size and colour
paraObj3 = doc.add_paragraph('This is a yet another paragraph but in bold, Calibri, size 11, red and +40% brightness', style = 'Body Text').bold = True

run4 = doc.add_paragraph().add_run()
#add a numbered list
paraObj4 = doc.add_paragraph('first item in ordered list')
paraObj4.style = 'List Number'

#add text to paragraph 1
paraObj1.add_run(' This text is being added to the second paragraph.')

#add a picture to your text of size 1 inch wide and 4 cm high
#doc.add_picture('zophie.png', width=docx.shared.Inches(1),height=docx.shared.Cm(4))
#if you use Inches or Cm lots then add:
#from docx.shared import Inches
#at the beginning and you can replace docx.shared.Inches with Inches
#similar for Cm

recordset = [record(1,101,'Spam'),record(2,42,'Eggs'),record(3,631,'Spam, spam, eggs and spam')]

#add a table
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
#add your items from a list of dictionaries to make it easy
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.qty)
    row_cells[1].text = str(item.id)
    row_cells[2].text = item.desc

#new page after
doc.add_page_break()

#save new work document as
doc.save('helloworld.docx')